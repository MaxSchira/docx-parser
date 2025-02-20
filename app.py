import json
from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import RGBColor, Pt
from io import BytesIO

app = Flask(__name__)

# ENDPOINT 1: DOCX-PARSING
@app.route('/parse-docx', methods=['POST'])
def parse_docx():
    file = request.files['file']
    doc = Document(file)
    extracted_text = []
    
    paragraph_counter = 0

    for para in doc.paragraphs:
        paragraph_counter += 1  
        paragraph_data = {
            "paragraph_id": paragraph_counter,
            "is_empty": len(para.text.strip()) == 0,
            "runs": []
        }

        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue  

            styles = {
                "font": run.font.name if run.font and run.font.name else "Calibri",
                "size": run.font.size.pt if run.font and run.font.size else 12,
                "bold": run.bold if run.bold else False,
                "italic": run.italic if run.italic else False,
                "underline": run.underline if run.underline else False,
                "color": f"#{run.font.color.rgb}" if run.font.color and isinstance(run.font.color.rgb, RGBColor) else "#000000"
            }

            paragraph_data["runs"].append({"text": text, "style": styles})

        extracted_text.append(paragraph_data)

    return jsonify({"extracted_data": extracted_text})

# ENDPOINT 2: TEXT ERSETZEN UND DOCX GENERIEREN
@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        # JSON aus Form-Data holen
        updated_speisekarte_json = request.form.get("updated_speisekarte", "[]")

        # JSON-Daten parsen und sicherstellen, dass es eine Liste ist
        try:
            updated_speisekarte = json.loads(updated_speisekarte_json)
            if not isinstance(updated_speisekarte, list):
                raise ValueError("updated_speisekarte ist keine Liste.")
        except json.JSONDecodeError:
            return jsonify({"error": "Fehler beim Parsen von updated_speisekarte"}), 400

        # Bestehendes DOCX-Dokument öffnen
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "Keine Datei empfangen"}), 400

        doc = Document(file)

        # Vorherigen Text entfernen
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = ""

        # Aktualisierte Speisekarte einfügen
        for item in updated_speisekarte:
            # Leere Absätze behandeln
            if item.get("is_empty", False):
                doc.add_paragraph("")  # Nur wenn nötig
                continue
    
            # Falls Struktur direkt "text" enthält (wie Weinkarte), konvertiere in runs-Format
            if "text" in item:
                item = {"runs": [{"text": item["text"], "style": item.get("style", {})}]}

            para = doc.add_paragraph()
            for run_data in item["runs"]:
                run = para.add_run(run_data["text"])
        
                # Formatierungen setzen
                style = run_data.get("style", {})
                run.bold = style.get("bold", False)
                run.italic = style.get("italic", False)
                run.underline = style.get("underline", False)
                run.font.name = style.get("font", "Futura Medium")
                run.font.size = Pt(style.get("size", 12))
        
                if "color" in style and style["color"] != "#000000":
                    rgb = style["color"].lstrip("#")
                    run.font.color.rgb = RGBColor(int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))

        # Datei speichern und zurücksenden
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        #debug log
        print("Received updated_speisekarte:", updated_speisekarte) 
        
        return send_file(output, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name="updated_menu.docx")

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=10000)
